# -*- coding: utf-8 -*-
"""
preserva_docx.py
----------------
Tradueix documents .docx preservant el format original run a run.

Correccions aplicades:
  - obte_runs_complets(): inclou runs dins d'hipervincles (w:hyperlink).
  - substitueix_text_paragraf(): nova estratègia que buida tots els runs
    (inclosos els d'hipervincles), restaura el format de forma condicional
    (només si el valor és explícit, no None) i estableix la llengua ca-ES.
  - neteja_traduccio(): versió robusta (Markdown complet + Unicode espuri).
  - Logging DEBUG per diagnosticar text brut vs. text net del model.
"""

import io
import logging
import re
import unicodedata

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph as DocxPara
from docx.text.run import Run

log = logging.getLogger(__name__)


# ── Helpers globals ────────────────────────────────────────────────────────────

# Detecta URLs per protegir-les abans de traduir (el model les trenca afegint espais)
URL_PATTERN = re.compile(
    r'https?://[^\s<>"\']+|www\.[^\s<>"\']+',
    re.IGNORECASE
)


def protegeix_urls(text: str) -> tuple[str, dict]:
    """Substitueix les URLs per marcadors i retorna el text modificat i el mapa.

    Usa el format URLTOKEN{n}XEND per evitar que el model el processe com
    a paraula normal (els guions baixos i els dobles guions sí que els modifica).
    """
    urls: dict = {}

    def reemplaca(m):
        key = f'URLTOKEN{len(urls)}XEND'
        urls[key] = m.group(0)
        return key

    text_protegit = URL_PATTERN.sub(reemplaca, text)
    return text_protegit, urls


def restaura_urls(text: str, urls: dict) -> str:
    """Restaura les URLs originals als marcadors."""
    for key, url in urls.items():
        text = text.replace(key, url)
    return text


def te_format_explicit(fmt: dict) -> bool:
    """Retorna True si el run tenia almenys un atribut de format definit explícitament."""
    return any([
        fmt['bold'] is not None,
        fmt['italic'] is not None,
        fmt['underline'] is not None,
        fmt['font_name'] is not None,
        fmt['font_size'] is not None,
        fmt['color'] is not None,
    ])


def distribueix_text_entre_runs(paraules: list, formats: list) -> list:
    """
    Distribueix les paraules proporcionalment entre els runs assegurant
    que no es perden espais entre paraules per arrodoniment.
    """
    total_original = sum(f['text_len'] for f in formats)
    total_paraules = len(paraules)
    fragments = []
    paraula_idx = 0

    for i, fmt in enumerate(formats):
        if i == len(formats) - 1:
            fragment = ' '.join(paraules[paraula_idx:])
        else:
            if total_original > 0:
                proporcio = fmt['text_len'] / total_original
                n_paraules = max(1, round(total_paraules * proporcio))
            else:
                n_paraules = 1
            fragment = ' '.join(paraules[paraula_idx:paraula_idx + n_paraules])
            paraula_idx = min(paraula_idx + n_paraules, total_paraules)
        fragments.append(fragment)

    return fragments


def corregeix_paraules_enganxades(text: str) -> str:
    """
    Detecta paraules enganxades (minúscula seguida directament de majúscula
    o lletres seguides de números sense espai) i afegeix l'espai que falta.
    Aplica NOMÉS als casos clars per evitar falsos positius.
    """
    # Cas: lletra minúscula enganxada amb majúscula (ex: "serveiEl" → "servei El")
    text = re.sub(r'([a-záéíóúàèìòùïüç])([A-ZÁÉÍÓÚÀÈÌÒÙÏÜÇ])', r'\1 \2', text)
    # Cas: número enganxat amb lletra (ex: "2023la" → "2023 la")
    text = re.sub(r'(\d)([A-ZÁÉÍÓÚÀÈÌÒÙÏÜÇa-záéíóúàèìòùïüç])', r'\1 \2', text)
    # Cas: lletra enganxada amb número (ex: "any2023" → "any 2023")
    # EXCEPCIÓ: no se separen URLs ni codis tècnics (ja protegits amb marcadors)
    text = re.sub(r'([a-záéíóúàèìòùïüç])(\d)', r'\1 \2', text)
    return text


def corregeix_apostrofacions(text: str) -> str:
    """
    Corregeix els espais incorrectes en apostrofacions catalanes/valencianes.
    Casos: l' , d' , m' , t' , s' , n' , c' , j' , qu'
    Exemples: "d' exemple" → "d'exemple", "l' ús" → "l'ús"
    """
    # Elimina espai entre apòstrof i paraula següent
    text = re.sub(r"([dlmtsncjq]u?)'(\s+)([^\s])", r"\1'\3", text, flags=re.IGNORECASE)

    # "de " + vocal/h → "d'" (prepositional elision)
    # NOTA: pot produir falsos positius en noms propis; descomenta si cal
    text = re.sub(
        r'\bde\s+([aeiouàèéíïòóúüh])',
        lambda m: "d'" + m.group(1),
        text,
        flags=re.IGNORECASE,
    )

    # Normalitza espais múltiples que puguen haver quedat
    text = re.sub(r' {2,}', ' ', text)

    return text.strip()


def neteja_traduccio(text: str) -> str:
    """Elimina tots els artefactes Markdown i caràcters espuris del model."""
    if not text:
        return text

    # Elimina negreta i cursiva Markdown (**text**, *text*, __text__)
    text = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'_{1,2}(.*?)_{1,2}', r'\1', text, flags=re.DOTALL)

    # Elimina capçaleres Markdown
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)

    # Elimina guions de llista al principi de línia
    text = re.sub(r'^[\-\–\—\•]\s+', '', text, flags=re.MULTILINE)

    # Elimina numeració al principi de línia AMB puntuació (1. text, 1) text)
    text = re.sub(r'^\d+[\.\)\-]\s+', '', text, flags=re.MULTILINE)

    # Elimina número sol al principi de línia SENSE puntuació (el cas "1 Segons")
    text = re.sub(r'^\d+\s+(?=[A-ZÁÉÍÓÚÀÈÌÒÙÏÜÇ·])', '', text, flags=re.MULTILINE)

    # Elimina número sol al principi de text (sense salts de línia)
    text = re.sub(r'^\d+\s+', '', text)

    # Elimina asteriscos solts que hagin quedat
    text = re.sub(r'\*+', '', text)

    # Normalitza espais de no-ruptura i caràcters de control Unicode
    text = re.sub(r'[\u00a0\u200b\u200c\u200d\ufeff]', ' ', text)
    text = re.sub(r' {2,}', ' ', text)

    # Elimina espais davant de puntuació
    text = re.sub(r' ([.,;:!?»\)\]])', r'\1', text)

    # Elimina línies buides múltiples
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


def _estableix_llengua_catalana(run) -> None:
    """
    Afegeix o actualitza l'atribut w:lang ca-ES al run de Word.
    Substitueix qualsevol element w:lang existent per evitar duplicats.
    """
    rPr = run._r.get_or_add_rPr()
    for lang in rPr.findall(qn('w:lang')):
        rPr.remove(lang)
    lang_elem = OxmlElement('w:lang')
    lang_elem.set(qn('w:val'), 'ca-ES')
    rPr.append(lang_elem)


def obte_runs_complets(para) -> list:
    """
    Retorna tots els runs del paràgraf, inclosos els dins d'hipervincles
    (w:hyperlink), que python-docx no inclou a para.runs per defecte.
    """
    runs = list(para.runs)
    for hipervinc in para._p.findall('.//' + qn('w:hyperlink')):
        for r_elem in hipervinc.findall(qn('w:r')):
            runs.append(Run(r_elem, para))
    return runs


def _aplica_format(run, fmt: dict) -> None:
    """
    Aplica el format guardat a un run DOCX.
    Escriu NOMÉS els valors explícits (no None) per evitar crear elements
    <w:rPr/> buits sobre runs que no tenien format propi.
    """
    if fmt['bold'] is not None:
        run.bold = fmt['bold']
    if fmt['italic'] is not None:
        run.italic = fmt['italic']
    if fmt['underline'] is not None:
        run.underline = fmt['underline']
    if fmt['font_name']:
        run.font.name = fmt['font_name']
    if fmt['font_size']:
        run.font.size = fmt['font_size']
    if fmt['color']:
        try:
            run.font.color.rgb = fmt['color']
        except Exception:
            pass


def substitueix_text_paragraf(para, text_traduit: str) -> None:
    """
    Substitueix el text del paràgraf distribuint el text traduït
    proporcionalment entre els runs originals i preservant el format de cada un.

    Estratègia:
      1. Captura el format de TOTS els runs (inclosos els d'hipervincles).
      2. Distribueix el text traduït amb distribueix_text_entre_runs().
      3. Aplica _aplica_format() NOMÉS als runs que tenien format explícit
         (te_format_explicit) per evitar afegir negreta/cursiva fantasma.
      4. Estableix la llengua ca-ES a tots els runs modificats.
    """
    runs_tots = obte_runs_complets(para)
    if not runs_tots:
        return

    # ── Captura el format de TOTS els runs ────────────────────────────────────
    formats = []
    for run in runs_tots:
        fmt = {
            'bold':      run.bold,
            'italic':    run.italic,
            'underline': run.underline,
            'font_name': run.font.name,
            'font_size': run.font.size,
        }
        try:
            fmt['color'] = run.font.color.rgb
        except Exception:
            fmt['color'] = None
        fmt['text_len'] = len(run.text)
        formats.append(fmt)

    total_original = sum(f['text_len'] for f in formats)

    # ── Cas simple: un sol run o tots els runs buits ───────────────────────────
    if total_original == 0 or len(runs_tots) == 1:
        runs_tots[0].text = text_traduit
        if te_format_explicit(formats[0]):
            _aplica_format(runs_tots[0], formats[0])
        _estableix_llengua_catalana(runs_tots[0])
        return

    # ── Buida el text de TOTS els runs (preserva l'estructura XML) ────────────
    for run in runs_tots:
        run.text = ''

    # ── Distribueix i aplica ───────────────────────────────────────────────────
    paraules = text_traduit.split()
    fragments = distribueix_text_entre_runs(paraules, formats)

    for run, fmt, fragment in zip(runs_tots, formats, fragments):
        run.text = fragment
        if te_format_explicit(fmt):
            _aplica_format(run, fmt)
        _estableix_llengua_catalana(run)


# ── Classe principal ───────────────────────────────────────────────────────────

class PreservadorDocx:
    def __init__(self, traductor):
        self.traductor = traductor

    def tradueix_segment(self, text: str) -> str:
        """
        Tradueix un segment de text aplicant tota la canonada de processament:
          1. Retorna sense canvis si el text és buit o és únicament una URL.
          2. Protegeix les URLs amb marcadors URLTOKEN{n}XEND.
          3. Tradueix amb el model.
          4. Neteja artefactes Markdown (neteja_traduccio).
          5. Restaura les URLs originals.
          6. Corregeix paraules enganxades (corregeix_paraules_enganxades).
          7. Corregeix apostrofacions incorrectes (corregeix_apostrofacions).
        """
        text = text.strip()
        if not text:
            return text
        # Si el segment és únicament una URL, no el tradueixis
        if URL_PATTERN.fullmatch(text):
            return text
        # Protegeix les URLs dins del text
        text_protegit, urls = protegeix_urls(text)
        # Traducció + neteja amb logging DEBUG per diagnosticar artefactes
        traduit_raw = self.traductor(text_protegit)
        log.debug("NETEJA IN:  %r", traduit_raw[:120])
        traduit = neteja_traduccio(traduit_raw)
        log.debug("NETEJA OUT: %r", traduit[:120])
        # Restaura les URLs originals
        traduit = restaura_urls(traduit, urls)
        # Postprocessament
        traduit = corregeix_paraules_enganxades(traduit)
        traduit = corregeix_apostrofacions(traduit)
        return traduit

    def tradueix_paragraf(self, para):
        """
        Tradueix el text d'un paràgraf delegant a tradueix_segment i
        substitueix_text_paragraf. Protegeix camps especials (TOC, etc.).
        """
        text = para.text.strip()
        runs_tots = obte_runs_complets(para)
        if not text or not runs_tots:
            return

        # Protegeix camps especials (TOC, camps calculats, etc.)
        for run in runs_tots:
            if run._r.find(qn('w:fldChar')) is not None:
                return
            if run._r.find(qn('w:instrText')) is not None:
                return

        traduit = self.tradueix_segment(text)
        substitueix_text_paragraf(para, traduit)

    def tradueix_taula(self, taula):
        for fila in taula.rows:
            for cel in fila.cells:
                for para in cel.paragraphs:
                    self.tradueix_paragraf(para)

    def tradueix_document(self, entrada, sortida=None):
        if isinstance(entrada, bytes):
            entrada = io.BytesIO(entrada)
        doc = Document(entrada)
        # Paràgrafs del cos principal
        for para in doc.paragraphs:
            self.tradueix_paragraf(para)
        # Taules
        for taula in doc.tables:
            self.tradueix_taula(taula)
        # Capçaleres i peus de pàgina
        for seccio in doc.sections:
            for part in [seccio.header, seccio.footer]:
                try:
                    for para in part.paragraphs:
                        self.tradueix_paragraf(para)
                except Exception:
                    pass
        # Quadres de text (textboxes)
        for txbx in doc.element.body.iter(qn('w:txbxContent')):
            for p_elem in txbx.iter(qn('w:p')):
                try:
                    self.tradueix_paragraf(DocxPara(p_elem, doc))
                except Exception:
                    pass
        buf = io.BytesIO()
        doc.save(buf)
        if sortida:
            doc.save(sortida)
        return buf.getvalue()
